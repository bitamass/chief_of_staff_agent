"""Repository-driven, deterministic execution of Chief of Staff skill contracts.

The runner reads SKILL.md, INPUTS.md, and OUTPUTS.md from the deployed GitHub
checkout.  OUTPUTS.md supplies the report structure; the selected initiative's
synthetic records supply the evidence.  No model API or production connection
is used.
"""

from __future__ import annotations

import re
from datetime import date, datetime
from io import BytesIO

from docx import Document


REQUIRED_DOCUMENTS = ("SKILL.md", "INPUTS.md", "OUTPUTS.md")


def _plain_heading(value: str) -> str:
    return re.sub(r"[*_`]", "", value).strip().rstrip(":")


def _headings(markdown: str) -> list[str]:
    """Return the most useful output section headings from OUTPUTS.md."""
    found: list[tuple[int, str]] = []
    for line in markdown.splitlines():
        match = re.match(r"^(#{2,4})\s+(.+?)\s*$", line)
        if match:
            found.append((len(match.group(1)), _plain_heading(match.group(2))))
    if not found:
        return ["Executive summary", "Evidence and findings", "Recommended next steps"]
    minimum = min(level for level, _ in found)
    sections = [heading for level, heading in found if level == minimum]
    return sections[:16]


def _bullets(markdown: str) -> list[str]:
    items = []
    for line in markdown.splitlines():
        match = re.match(r"^\s*[-*]\s+(.+?)\s*$", line)
        if match:
            item = re.sub(r"[*_`]", "", match.group(1)).strip()
            if item and item not in items:
                items.append(item)
    return items


def _fmt(value) -> str:
    if isinstance(value, (date, datetime)):
        return value.strftime("%b %d, %Y")
    return str(value)


def _source(label: str, text: str) -> dict:
    return {"text": text, "source": label}


def _person(context, person_id: str) -> str:
    person = context["participants"][person_id]
    return f"{person['name']} ({person['role']})"


def _initiative_evidence(context, scenario_name: str) -> dict:
    """Normalize the selected initiative into reusable evidence families."""
    from synthetic_context import build_initiative_scenarios, initiative_continuity

    scenario = build_initiative_scenarios(context)[scenario_name]
    initiative_id = scenario["initiative_id"]
    initiative = context["initiatives"][initiative_id]
    continuity = initiative_continuity(context, initiative_id)
    charter = context["charters"][initiative["charter_id"]]
    deliverables = [d for d in context["deliverables"].values() if d["initiative_id"] == initiative_id]

    participants = []
    for person in continuity["participants"]:
        participants.append(_source("Synthetic participant register", f"{person['name']} — {person['role']}"))

    meetings = []
    for item in reversed(continuity["prior_meetings"]):
        meetings.append(_source(item["title"], f"Prior · {_fmt(item['date'])}: {item['summary']}"))
    for item in continuity["current_meetings"]:
        meetings.append(_source(item["title"], f"Today · {_fmt(item['date'])}: {item['summary']}"))
    for item in continuity["upcoming_meetings"]:
        meetings.append(_source(item["title"], f"Upcoming · {_fmt(item['date'])}: {item['summary']}"))

    decisions = []
    for item in reversed(continuity["prior_decisions"]):
        source_meeting = context["meetings"][item["meeting_id"]]["title"]
        decisions.append(_source(source_meeting, f"Decided {_fmt(item['date'])}: {item['decision']}"))

    upcoming_decisions = []
    for item in continuity["upcoming_decisions"]:
        source_meeting = context["meetings"][item["meeting_id"]]["title"]
        upcoming_decisions.append(_source(source_meeting, f"Decision due {_fmt(item['date'])}: {item['decision']}"))

    decision_summary = [
        _source(
            "Synthetic decision register",
            f"Decision record contains {len(decisions)} prior decision(s) and {len(upcoming_decisions)} upcoming decision(s) for {initiative['name']}.",
        )
    ]
    if decisions:
        most_recent = continuity["prior_decisions"][0]
        decision_summary.append(
            _source(
                context["meetings"][most_recent["meeting_id"]]["title"],
                f"Most recent decision · {_fmt(most_recent['date'])}: {most_recent['decision']}",
            )
        )
    if continuity["upcoming_decisions"]:
        next_decision = continuity["upcoming_decisions"][0]
        decision_summary.append(
            _source(
                context["meetings"][next_decision["meeting_id"]]["title"],
                f"Next decision required · {_fmt(next_decision['date'])}: {next_decision['decision']}",
            )
        )

    decision_log_entries = []
    for item in reversed(continuity["prior_decisions"]):
        source_meeting = context["meetings"][item["meeting_id"]]
        decision_log_entries.append(
            _source(
                source_meeting["title"],
                f"{item['id']} | Decided {_fmt(item['date'])} | {item['decision']} | Initiative: {initiative['name']} | Meeting: {source_meeting['title']}.",
            )
        )

    complete_records = sum(
        bool(item.get("id") and item.get("date") and item.get("decision") and item.get("meeting_id") and item.get("initiative_id"))
        for item in continuity["prior_decisions"]
    )
    record_completeness = [
        _source(
            "Synthetic decision register",
            f"Required demo fields are present for {complete_records} of {len(decisions)} prior decision record(s): ID, date, decision statement, initiative, and source meeting.",
        ),
        _source(
            "Completeness assessment",
            "Not represented in the synthetic decision records: named decision owner, decision rationale, alternatives considered, approval authority, and supersession status.",
        ),
        _source(
            "Completeness assessment",
            f"{len(upcoming_decisions)} upcoming decision(s) are documented separately and should not be reported as completed decisions.",
        ),
    ]

    decision_human_review = [
        _source("Human-review control", "Confirm each decision statement and date against the cited meeting record."),
        _source("Human-review control", "Add or validate the decision owner, approval authority, rationale, and alternatives considered before formal use."),
        _source("Human-review control", "Confirm that upcoming decisions are clearly separated from decisions already made."),
        _source("Human-review control", "Resolve contradictions or superseded decisions and approve the log before distribution."),
    ]

    actions = []
    for item in continuity["all_actions"]:
        owner = context["participants"][item["owner"]]
        actions.append(
            _source(
                context["meetings"][item["source_meeting_id"]]["title"],
                f"{item['status']} · Due {_fmt(item['due_date'])}: {item['action']} Owner: {owner['name']} ({owner['role']}).",
            )
        )

    risks = []
    risk_terms = ("risk", "security", "privacy", "dependency", "constraint", "incomplete", "delay", "control", "readiness")
    for meeting in meetings:
        if any(term in meeting["text"].lower() for term in risk_terms):
            risks.append(meeting)
    for action in actions:
        if any(term in action["text"].lower() for term in ("blocked", "not started", "dependency", "security", "risk")):
            risks.append(action)
    if not risks:
        risks.append(_source("Synthetic charter", "No explicit risk statement was found; validate dependencies, ownership, timing, and evidence before use."))

    deliverable_items = [
        _source("Synthetic deliverable register", f"{d['name']} — {d['status']}; due {_fmt(d['due_date'])}.")
        for d in sorted(deliverables, key=lambda x: x["due_date"])
    ]
    objectives = [
        _source("Synthetic initiative register", f"Initiative: {initiative['name']}. Status: {initiative['status']}."),
        _source("Synthetic initiative register", f"Objective: {initiative['objective']}"),
        _source("Synthetic charter", f"Scope: {charter['scope']}"),
        _source("Synthetic charter", f"Out of scope: {charter['out_of_scope']}"),
        _source("Synthetic charter", "Success measures: " + ", ".join(charter["success_measures"])),
    ]

    summary = [
        _source("Synthetic initiative register", f"{initiative['name']} is {initiative['status'].lower()} and supports: {initiative['objective']}"),
        _source("Synthetic calendar", f"Most relevant meeting: {scenario['title']} on {_fmt(scenario['date'])}."),
        _source("Synthetic continuity record", f"{len(decisions)} prior decisions, {len(actions)} actions, and {len(upcoming_decisions)} upcoming decisions are linked to this initiative."),
    ]

    questions = [_source("Synthetic meeting context", question) for question in scenario["questions"]]
    recommendations = [_source("Synthetic action register", step) for step in scenario["next_steps"]]
    gaps = [
        _source("Demo limitation", "Only synthetic records in the five-week demonstration window were evaluated."),
        _source("Demo limitation", "No production-system validation, external intelligence, or autonomous action was performed."),
    ]

    priorities = []
    for index, item in enumerate(sorted(deliverables, key=lambda value: value["due_date"]), start=1):
        priorities.append(
            _source(
                "Synthetic deliverable register",
                f"Priority {index}: {item['name']} — {item['status']}; owned by {_person(context, item['owner'])}; due {_fmt(item['due_date'])}.",
            )
        )
    for item in continuity["upcoming_decisions"][:3]:
        priorities.append(
            _source(
                context["meetings"][item["meeting_id"]]["title"],
                f"Decision priority: {item['decision']} Required by {_fmt(item['date'])}.",
            )
        )

    readiness = []
    for item in sorted(deliverables, key=lambda value: value["due_date"]):
        status = "Ready for review" if item["status"].lower() in {"complete", "draft complete"} else "Preparation required"
        readiness.append(
            _source(
                "Synthetic deliverable register",
                f"{status}: {item['name']} is {item['status'].lower()}, owned by {_person(context, item['owner'])}, and due {_fmt(item['due_date'])}.",
            )
        )
    readiness.extend(risks[:3])

    schedule = []
    for item in continuity["upcoming_meetings"][:5]:
        schedule.append(_source(item["title"], f"{_fmt(item['date'])}: {item['title']} — {item['summary']}"))
    for item in sorted(deliverables, key=lambda value: value["due_date"]):
        schedule.append(_source("Synthetic deliverable register", f"Due {_fmt(item['due_date'])}: {item['name']} — {item['status']}."))

    ownership = participants[:]
    for item in continuity["all_actions"]:
        ownership.append(
            _source(
                "Synthetic action register",
                f"{_person(context, item['owner'])} owns: {item['action']} Status: {item['status']}; due {_fmt(item['due_date'])}.",
            )
        )

    raci_rows = []
    initiative_owner = _person(context, initiative["owner"])
    for item in sorted(deliverables, key=lambda value: value["due_date"]):
        responsible = _person(context, item["owner"])
        raci_rows.append(
            _source(
                "Synthetic deliverable and initiative registers",
                f"Deliverable: {item['name']} | R: {responsible} | A (provisional): {initiative_owner} | C: Not documented | I: Not documented | Due: {_fmt(item['due_date'])} | Validation: Formal authority and consultation assignments require confirmation.",
            )
        )
    for item in continuity["all_actions"]:
        responsible = _person(context, item["owner"])
        raci_rows.append(
            _source(
                context["meetings"][item["source_meeting_id"]]["title"],
                f"Action: {item['action']} | R: {responsible} | A (provisional): {initiative_owner} | C: Not documented | I: Not documented | Due: {_fmt(item['due_date'])} | Status: {item['status']} | Validation: Provisional assignment.",
            )
        )
    for item in continuity["upcoming_decisions"]:
        raci_rows.append(
            _source(
                context["meetings"][item["meeting_id"]]["title"],
                f"Decision: {item['decision']} | R: Not documented | A: Not documented | C: Meeting participants are candidates only | I: Not documented | Due: {_fmt(item['date'])} | Validation: Decision authority must be confirmed.",
            )
        )

    responsible_counts = {}
    for item in list(deliverables) + list(continuity["all_actions"]):
        responsible_counts[item["owner"]] = responsible_counts.get(item["owner"], 0) + 1
    role_summaries = []
    for person_id in initiative["participant_ids"]:
        assignments = responsible_counts.get(person_id, 0)
        accountable = "Provisional initiative accountability" if person_id == initiative["owner"] else "No accountable assignment documented"
        role_summaries.append(
            _source(
                "Synthetic participant, action, and deliverable registers",
                f"{_person(context, person_id)} | Responsible assignments: {assignments} | Accountability: {accountable} | Approval authority: Not documented | Consultation/information duties: Not documented | Backup: Not documented.",
            )
        )

    raci_gaps = [
        _source("RACI validation", f"{len(continuity['upcoming_decisions'])} upcoming decision(s) lack a documented Responsible role and accountable decision authority."),
        _source("RACI validation", "Consulted and Informed assignments are not explicitly documented for actions or deliverables."),
        _source("RACI validation", "The initiative owner is used only as a provisional Accountable role; formal delegation and approval authority are not represented."),
        _source("RACI validation", "Acceptance owners, escalation paths, delegation arrangements, and backup coverage are not represented in the synthetic records."),
    ]

    same_ra = [
        item["name"] if "name" in item else item["action"]
        for item in list(deliverables) + list(continuity["all_actions"])
        if item["owner"] == initiative["owner"]
    ]
    raci_conflicts = [
        _source(
            "RACI validation",
            ("Potential R/A concentration requiring review: " + "; ".join(same_ra))
            if same_ra
            else "No row has the same documented Responsible owner and provisional Accountable owner in the synthetic records.",
        ),
        _source("RACI validation", "Multiple accountable owners are not documented; this cannot be confirmed without formal governance and delegation records."),
        _source("RACI validation", "Consultation overload, overlapping team assignments, and authority-resource conflicts cannot be assessed because C/I roles, capacity, and delegation data are absent."),
    ]

    raci_workload = [
        _source("Synthetic assignment register", f"{_person(context, person_id)} has {count} Responsible assignment(s) across documented actions and deliverables.")
        for person_id, count in sorted(responsible_counts.items(), key=lambda pair: (-pair[1], pair[0]))
    ]
    raci_workload.append(_source("Workload limitation", "Availability, percentage allocation, scheduling conflicts, and backup capacity are not represented; workload risk cannot be fully scored."))

    raci_governance = [
        _source("Draft recommendation — human approval required", "Confirm one accountable authority for every deliverable, action, and decision."),
        _source("Draft recommendation — human approval required", "Assign Consulted and Informed roles based on required expertise, governance, and communication needs."),
        _source("Draft recommendation — human approval required", "Document escalation, delegation, backup, acceptance, and approval authority before operational use."),
        _source("Draft recommendation — human approval required", f"Review and approve the RACI at {scenario['title']} on {_fmt(scenario['date'])}, then reassess when scope, ownership, or governance changes."),
    ]

    raci_executive_summary = [
        _source("RACI assessment", f"{len(raci_rows)} work or decision row(s) were assessed for {initiative['name']}."),
        _source("RACI assessment", f"Documented Responsible assignments exist for actions and deliverables; {len(continuity['upcoming_decisions'])} upcoming decision(s) still lack explicit decision authority."),
        _source("RACI assessment", "Material gaps remain in formal accountability, consultation, informed parties, approval authority, escalation, and backup coverage."),
        _source("RACI assessment", "Do not treat provisional assignments as an approved RACI until accountable leaders and governance partners validate them."),
    ]

    outcomes = [
        _source("Synthetic initiative register", f"Advance the initiative objective: {initiative['objective']}"),
        _source("Synthetic charter", f"Weekly success should be assessed using: {', '.join(charter['success_measures'])}."),
    ]
    outcomes.extend(
        _source("Synthetic deliverable register", f"Move {item['name']} from {item['status'].lower()} to its next documented review state by {_fmt(item['due_date'])}.")
        for item in sorted(deliverables, key=lambda value: value["due_date"])
    )

    success_measures = [
        _source("Synthetic charter", f"Measure: {measure}. Confirm the definition, baseline, target, and evidence owner before reporting success.")
        for measure in charter["success_measures"]
    ]

    # The synthetic dataset does not contain approved institutional OKRs. For
    # this skill, the initiative objective and charter success measures are
    # evaluated transparently as demonstration proxies rather than represented
    # as approved OKRs.
    okr_summary = [
        _source(
            "Synthetic initiative and charter records",
            f"One initiative-level objective and {len(charter['success_measures'])} candidate success measure(s) were available as OKR proxies for {initiative['name']}.",
        ),
        _source(
            "OKR evidence assessment",
            "Alignment to approved enterprise strategic priorities is unclear because no approved strategic-priority or organization-level OKR record is included in the synthetic dataset.",
        ),
        _source(
            "OKR evidence assessment",
            "The candidate measures name intended outcomes but lack documented baselines, targets, current values, measurement periods, and data owners.",
        ),
    ]

    okr_inventory = [
        _source(
            "Synthetic initiative register",
            f"Objective proxy {initiative_id}: {initiative['objective']} Level: Initiative; owner: {_person(context, initiative['owner'])}; status: {initiative['status']}; planning period: five-week demonstration window.",
        )
    ]
    okr_inventory.extend(
        _source(
            "Synthetic charter",
            f"Candidate key-result proxy KR-{index}: {measure}. Baseline, target, current value, measurement period, frequency, and accountable data owner are not documented.",
        )
        for index, measure in enumerate(charter["success_measures"], start=1)
    )

    objective_quality = [
        _source("Synthetic initiative register", f"Strategic relevance: Unverified — no approved strategic-priority record is available to validate alignment for '{initiative['objective']}'."),
        _source("Objective-quality assessment", "Clarity and outcome orientation: Partially demonstrated — the objective names the intended operational improvement but does not define a quantified end state."),
        _source("Objective-quality assessment", f"Ownership: Represented by {_person(context, initiative['owner'])}; confirm formal accountability before use."),
        _source("Objective-quality assessment", "Time-bound focus: Needs clarification — the demonstration window is not an approved objective period or deadline."),
        _source("Objective-quality assessment", "Overall classification: Needs clarification before it can be treated as an approved OKR objective."),
    ]

    key_result_quality = []
    for index, measure in enumerate(charter["success_measures"], start=1):
        key_result_quality.append(
            _source(
                "Synthetic charter",
                f"KR-{index} proxy · {measure}: classification cannot be confirmed as outcome, output, activity, or milestone. Missing baseline, target, current value, unit, period, frequency, data source, and accountable owner; confidence is low.",
            )
        )

    vertical_alignment = [
        _source("Synthetic initiative register", f"Initiative objective: {initiative['objective']}"),
        _source("Synthetic charter", f"Candidate measures supporting the objective: {', '.join(charter['success_measures'])}."),
        _source("Synthetic deliverable register", "Supporting work: " + "; ".join(item["name"] for item in deliverables) + "."),
        _source("Alignment limitation", "Enterprise, division, department, and team objectives are not represented; vertical alignment above the initiative level cannot be validated."),
    ]

    horizontal_alignment = [
        _source("Synthetic participant register", "Cross-functional roles represented: " + "; ".join(person["text"] for person in participants) + "."),
        _source("Synthetic continuity record", f"Coordination demand: {len(continuity['upcoming_meetings'])} upcoming meeting(s) and {len(open_actions) if 'open_actions' in locals() else len(actions)} documented action record(s)."),
        _source("Alignment limitation", "Peer-team OKRs, shared targets, metric ownership, resource conflicts, and duplicate objectives are not represented in the synthetic dataset."),
    ]

    okr_coverage_gaps = [
        _source("Coverage assessment", "Approved strategic priorities and parent objectives are missing, so strategic coverage cannot be scored."),
        _source("Coverage assessment", "The objective proxy lacks formally defined and approved measurable key results."),
        _source("Coverage assessment", "Candidate measures lack baselines, targets, current values, measurement periods, data sources, and measurement owners."),
        _source("Coverage assessment", "Initiative deliverables are documented, but their quantitative contribution to candidate key results is not."),
    ]

    okr_revisions = [
        _source("Draft recommendation — owner approval required", f"Clarify the objective with an approved outcome, target population or process, quantified end state, and deadline: {initiative['objective']}"),
        _source("Draft recommendation — owner approval required", "Convert each charter success measure into a defined key result with a baseline, target, current value, unit, period, data source, frequency, and accountable owner."),
        _source("Draft recommendation — owner approval required", "Map the initiative objective to an approved strategic priority and document the evidence and confidence for that relationship."),
        _source("Draft recommendation — owner approval required", "Document how each deliverable contributes to a key result and how that contribution will be measured."),
    ]

    okr_progress = [
        _source("Synthetic initiative register", f"Current initiative status: {initiative['status']}. This status is not evidence of OKR outcome achievement."),
        _source("Synthetic deliverable register", f"Delivery activity: {len(deliverables)} deliverable(s) are represented; activity completion must not be substituted for key-result performance."),
        _source("Progress limitation", "Progress and likelihood of achievement cannot be calculated because baseline, target, current-value, and measurement-period records are absent."),
        _source("Synthetic calendar", f"Next review opportunity: {scenario['title']} on {_fmt(scenario['date'])}."),
    ]

    okr_human_review = [
        _source("Human-review control", "The strategy owner must confirm the approved strategic priority and objective hierarchy."),
        _source("Human-review control", "Objective and key-result owners must approve wording, ownership, baselines, targets, periods, and data sources."),
        _source("Human-review control", "Metric owners must validate current values, measurement frequency, and evidence quality before progress is reported."),
        _source("Human-review control", "Draft mappings and revisions must not alter an official OKR or performance record without authorization."),
    ]

    capacity = []
    open_actions = [item for item in continuity["all_actions"] if item["status"].lower() != "complete"]
    owner_counts = {}
    for item in open_actions:
        owner_counts[item["owner"]] = owner_counts.get(item["owner"], 0) + 1
    for owner_id, count in sorted(owner_counts.items(), key=lambda value: (-value[1], value[0])):
        capacity.append(_source("Synthetic action register", f"{_person(context, owner_id)} has {count} open action(s) in this initiative during the five-week window."))
    capacity.append(
        _source(
            "Synthetic calendar",
            f"The initiative has {len(continuity['upcoming_meetings'])} upcoming meeting(s), {len(open_actions)} open action(s), and {len(continuity['upcoming_decisions'])} upcoming decision(s) in the demonstration window.",
        )
    )

    tradeoffs = []
    for item in continuity["upcoming_decisions"][:4]:
        tradeoffs.append(
            _source(
                context["meetings"][item["meeting_id"]]["title"],
                f"Trade-off for leadership: {item['decision']} Evaluate value, timing, risk, capacity, and the consequence of delay before selecting a direction.",
            )
        )

    approvals = upcoming_decisions or [
        _source("Synthetic governance record", "No formal upcoming approval was found for this initiative in the five-week window; confirm whether an approval is required.")
    ]

    status = [
        _source("Synthetic initiative register", f"Initiative status: {initiative['status']} — {initiative['name']}."),
        _source("Synthetic continuity record", f"Operating record: {len(meetings)} linked meetings, {len(decisions)} prior decisions, {len(actions)} actions, and {len(upcoming_decisions)} upcoming decisions."),
    ] + deliverable_items

    priority_rationale = []
    for item in sorted(deliverables, key=lambda value: value["due_date"]):
        priority_rationale.append(
            _source(
                "Synthetic charter and deliverable register",
                f"Prioritize {item['name']} because it supports '{initiative['objective']}', is currently {item['status'].lower()}, and is due {_fmt(item['due_date'])}. Validate effort and consequence before changing its rank.",
            )
        )

    dependency_map = []
    for item in continuity["all_actions"]:
        review = context["meetings"].get(item.get("review_meeting_id"), {})
        review_text = f" Review is scheduled at {review['title']} on {_fmt(review['date'])}." if review else ""
        dependency_map.append(
            _source(
                context["meetings"][item["source_meeting_id"]]["title"],
                f"Dependency: {item['action']} Owner: {_person(context, item['owner'])}; due {_fmt(item['due_date'])}.{review_text}",
            )
        )
    dependency_map.extend(risks[:3])

    return {
        "summary": summary,
        "overview": summary + objectives[:2],
        "objective": objectives,
        "alignment": objectives + deliverable_items,
        "participants": participants,
        "stakeholders": participants,
        "meetings": meetings,
        "timeline": meetings,
        "history": meetings + decisions,
        "decisions": decisions,
        "decision_summary": decision_summary,
        "decision_log_entries": decision_log_entries,
        "record_completeness": record_completeness,
        "decision_human_review": decision_human_review,
        "upcoming_decisions": upcoming_decisions,
        "actions": actions,
        "commitments": actions,
        "deliverables": deliverable_items,
        "priorities": priorities,
        "priority_rationale": priority_rationale,
        "outcomes": outcomes,
        "success_measures": success_measures,
        "okr_summary": okr_summary,
        "okr_inventory": okr_inventory,
        "objective_quality": objective_quality,
        "key_result_quality": key_result_quality,
        "vertical_alignment": vertical_alignment,
        "horizontal_alignment": horizontal_alignment,
        "okr_coverage_gaps": okr_coverage_gaps,
        "okr_revisions": okr_revisions,
        "okr_progress": okr_progress,
        "okr_human_review": okr_human_review,
        "readiness": readiness,
        "schedule": schedule,
        "ownership": ownership,
        "raci_matrix": raci_rows,
        "raci_role_summaries": role_summaries,
        "raci_gaps": raci_gaps,
        "raci_conflicts": raci_conflicts,
        "raci_workload": raci_workload,
        "raci_governance": raci_governance,
        "raci_executive_summary": raci_executive_summary,
        "capacity": capacity,
        "tradeoffs": tradeoffs,
        "dependency_map": dependency_map,
        "approvals": approvals,
        "status": status,
        "risks": risks,
        "issues": risks,
        "dependencies": risks,
        "questions": questions,
        "recommendations": recommendations,
        "next_steps": recommendations,
        "gaps": gaps,
        "confidence": gaps,
        "sources": [
            _source("Repository", "SKILL.md — workflow and guardrails"),
            _source("Repository", "INPUTS.md — required evidence contract"),
            _source("Repository", "OUTPUTS.md — report structure"),
            _source("Synthetic dataset", "Initiative, calendar, meeting, decision, action, charter, participant, and deliverable records"),
        ],
        "meeting": [_source("Synthetic calendar", f"{scenario['title']} · {_fmt(scenario['date'])} · {scenario['attendees']}")],
        "charter": objectives[2:],
        "governance": gaps,
    }


def _matches_heading(value: str, phrase: str) -> bool:
    """Match complete heading words/phrases, not accidental substrings.

    This prevents, for example, ``source`` from matching ``resource`` and
    ``action`` from matching ``satisfaction``.
    """
    pattern = r"(?<![a-z0-9])" + re.escape(phrase.lower()).replace(r"\ ", r"[\s-]+") + r"(?![a-z0-9])"
    return bool(re.search(pattern, value.lower()))


def _first(values: list[dict], fallback: str, source: str = "Demo assessment") -> dict:
    return values[0] if values else _source(source, fallback)


def _skill_specific_evidence(skill_name: str, heading: str, evidence: dict) -> list[dict] | None:
    """Return differentiated evidence for high-value analytical skills."""
    skill = skill_name.lower()
    value = heading.lower()
    summary = evidence["summary"]
    objective = _first(evidence["objective"], "No initiative objective is represented.")
    risks = evidence["risks"]
    actions = evidence["actions"]
    decisions = evidence["decisions"]
    upcoming = evidence["upcoming_decisions"]
    deliverables = evidence["deliverables"]

    if skill == "executive dashboard generation":
        most_recent_decision = _first(decisions, "No completed decision is represented.")
        next_decision = _first(upcoming, "No upcoming executive decision is represented.")
        next_action = _first(actions, "No open action is represented.")
        next_milestone = _first(evidence["schedule"], "No upcoming milestone is represented.")
        principal_risk = _first(risks, "No material risk is explicitly represented; owner validation is required.")
        mapping = {
            "dashboard specification": [
                _source("Dashboard design", "Purpose: give executive leaders a concise exception-oriented view of initiative status, decisions, actions, risks, and upcoming milestones."),
                _source("Dashboard design", "Audience: executive sponsor and initiative leadership. Scope: selected synthetic initiative. Cutoff: current demonstration date. Proposed refresh: before each executive review."),
                _source("Dashboard limitation", "No approved KPI definitions, targets, thresholds, financial actuals, access classification, or prior-period dashboard are represented; display those as gaps rather than invented values."),
            ],
            "executive summary panel": [summary[0], most_recent_decision, principal_risk, next_decision, next_milestone],
            "key performance indicators": [
                _source("Synthetic initiative register", "KPI candidate — Initiative status | Current value: documented initiative status | Target and threshold: not represented | Owner validation required."),
                _source("Synthetic action register", f"KPI candidate — Open commitments | Current evidence: {len(actions)} linked action(s) | Target and previous period: not represented."),
                _source("Synthetic decision register", f"KPI candidate — Pending executive decisions | Current evidence: {len(upcoming)} | Target and tolerance: not represented."),
                _source("Data-quality control", "Do not display calculated performance, trend arrows, or red/amber/green status until definitions, targets, periods, owners, and source timestamps are approved."),
            ],
            "strategic-priority view": [objective] + evidence["priorities"][:4],
            "initiative and portfolio view": summary[:3] + deliverables[:4],
            "financial and resource view": evidence["capacity"][:4] + [_source("Financial-data gap", "Approved budget, actual spending, forecast, benefits realized, staffing plan, and vendor commitments are not represented in the synthetic dataset.")],
            "risk and issue view": risks[:5],
            "decisions and actions view": [next_decision, next_action] + decisions[:3] + actions[:3],
            "milestone and forward-look view": evidence["schedule"][:6],
            "exception and attention panel": [principal_risk, next_decision, next_action, _source("Executive attention", "Validate missing ownership, thresholds, financial evidence, and stale or conflicting data before distribution.")],
            "drill-down and source view": evidence["sources"] + [_source("Traceability requirement", "Every future metric should include definition, calculation, owner, reporting period, last refresh, underlying record, and limitation.")],
            "data-quality panel": evidence["gaps"][:3] + [_source("Data-quality assessment", "Completeness is sufficient for a synthetic workflow demonstration but insufficient for official operational, financial, clinical, audit, or compliance reporting.")],
        }
        return mapping.get(value)

    if skill == "align deliverables with strategic goals":
        score_rows = []
        for index, item in enumerate(deliverables, start=1):
            alignment = max(2, 5 - index + 1)
            urgency = max(2, 5 - index)
            overall = round((alignment * .45) + (urgency * .25) + (4 * .2) + (3 * .1), 1)
            disposition = "Accelerate" if overall >= 4.2 else "Keep" if overall >= 3.5 else "Revise"
            score_rows.append(_source(item["source"], f"{item['text']} | Strategic alignment {alignment}/5 | Business contribution 4/5 | Expected value 4/5 | Urgency {urgency}/5 | Dependency importance 3/5 | Feasibility 3/5 | Weighted score {overall}/5 | Recommendation: {disposition}. Scores are demonstration assessments requiring owner validation."))
        mapping = {
            "executive summary": [summary[0], _source("Alignment assessment", f"{len(deliverables)} deliverable(s) were assessed against the documented initiative objective as a proxy; no approved enterprise strategic-goal record is available."), _source("Human-review control", "Do not change priorities or resources until the proxy goal, weights, and scores are approved.")],
            "strategic-goal evaluation": [objective, _source("Strategic-goal limitation", "The initiative objective is used as a transparent proxy because approved enterprise goals and top business priorities are not represented."), _source("Coverage assessment", f"{len(deliverables)} deliverable(s) provide partial coverage; enterprise-level coverage cannot be determined.")],
            "deliverable alignment matrix": score_rows,
            "prioritized deliverable list": [
                _source(item["source"], f"Rank {index}: {item['text'].split(' | ')[0]} | {item['text'].split('Recommendation: ')[-1]}")
                for index, item in enumerate(
                    sorted(score_rows, key=lambda row: float(re.search(r"Weighted score ([0-9.]+)", row["text"]).group(1)), reverse=True),
                    start=1,
                )
            ],
            "strategic coverage gaps": [_source("Coverage assessment", "Approved enterprise strategic goals, priority weights, targets, and measurement owners are missing."), _source("Coverage assessment", "Deliverable contribution is qualitatively supported but quantitative outcome contribution is not documented."), _source("Coverage assessment", "Validate resource capacity, dependencies, and completion criteria before final prioritization.")],
            "executive decisions and recommendations": (upcoming[:3] or evidence["recommendations"][:3]) + [_source("Draft recommendation", "Approve or revise the scoring criteria, weights, and recommended dispositions; assign owners for unresolved evidence gaps.")],
            "human-review requirement": [_source("Human-review control", "An authorized strategy and initiative owner must validate the proxy goal, scores, weights, rationale, and Keep/Accelerate/Revise recommendations before any change.")],
        }
        return mapping.get(value)

    if skill == "trade off analysis":
        alternatives = [
            _source("Demonstration option analysis", "Option A — proceed with the limited current scope. Gain: earlier learning and time to value. Sacrifice: incomplete assurance and possible rework. Reversibility: medium. Confidence: medium."),
            _source("Demonstration option analysis", "Option B — pause for missing cost, security, ownership, and readiness evidence. Gain: stronger assurance. Sacrifice: delay and lost learning time. Reversibility: high. Confidence: medium."),
            _source("Demonstration option analysis", "Option C — defer the initiative. Gain: preserves near-term capacity. Sacrifice: delays the documented objective and expected value. Reversibility: high. Confidence: low."),
        ]
        mapping = {
            "executive trade-off summary": (evidence["tradeoffs"][:3] or upcoming[:2]) + [_source("Trade-off synthesis", "Leadership is balancing speed and learning against assurance, capacity, and the cost of avoidable rework.")],
            "decision framing": (upcoming[:2] or decisions[:2]) + [objective, _source("Decision framing", "Success should be judged against the charter measures, constraints, available capacity, and consequence of delay.")],
            "alternative profiles": alternatives,
            "trade-off comparison matrix": [_source("Demonstration comparison", "Option A | Alignment 4/5 | Time to value 5/5 | Assurance 2/5 | Capacity demand 4/5 | Reversibility 3/5"), _source("Demonstration comparison", "Option B | Alignment 4/5 | Time to value 2/5 | Assurance 5/5 | Capacity demand 3/5 | Reversibility 5/5"), _source("Demonstration comparison", "Option C | Alignment 1/5 | Time to value 1/5 | Assurance 4/5 | Capacity demand 1/5 | Reversibility 5/5")],
            "gains and sacrifices": [_source("Trade-off synthesis", "Option A gains early learning but sacrifices assurance and may create rework."), _source("Trade-off synthesis", "Option B gains assurance and evidence quality but sacrifices speed and delays value."), _source("Trade-off synthesis", "Option C preserves near-term capacity but sacrifices momentum, learning, and progress toward the objective.")],
            "resource and opportunity-cost analysis": evidence["capacity"][:4] + [_source("Opportunity-cost assessment", "The same leadership, architecture, finance, clinical, and security capacity cannot simultaneously support all alternatives; validate effort and funding before selection.")],
            "risk and uncertainty analysis": risks[:4] + [_source("Uncertainty assessment", "Cost, benefit timing, stakeholder burden, and control effectiveness are not fully quantified; reconsider if a key dependency fails.")],
            "stakeholder impact analysis": evidence["participants"][:5],
            "sensitivity analysis": [_source("Sensitivity assessment", "Option A becomes less attractive if cost, implementation time, or control gaps increase."), _source("Sensitivity assessment", "Option B becomes more attractive when assurance requirements or uncertainty increase, but less attractive when delay costs rise."), _source("Sensitivity assessment", "The choice is most sensitive to time-to-value, readiness, available capacity, and the consequence of delay.")],
            "leadership judgment required": [_source("Leadership judgment", "Choose the acceptable balance between speed and assurance, near-term cost and long-term capability, and central control and local flexibility.")],
            "information gaps and confidence": evidence["gaps"][:3] + [_source("Confidence assessment", "Overall confidence is medium-low until comparable cost, benefit, resource, control, and stakeholder evidence is validated.")],
            "human-review requirement": [_source("Human-review control", "Authorized leadership must validate the alternatives, assumptions, comparison dimensions, risks, and opportunity costs before committing resources.")],
        }
        return mapping.get(value)

    if skill == "risk scoring":
        risk = _first(risks, "No explicit risk was documented; specialist validation is required.")
        mapping = {
            "risk-scoring record": [risk, _source("Demonstration scoring method", "Risk-01 | Operational and assurance readiness | Five-point probability × five-point impact matrix | Assessment owner: initiative owner | Confidence: low-to-medium.")],
            "inherent risk assessment": [_source("Demonstration risk assessment", "Probability 4/5; impact 4/5; inherent rating 16/25 (High). Rationale: unresolved dependencies could delay delivery and weaken assurance. Velocity: medium; detectability: medium; confidence: low-to-medium."), risk],
            "control assessment": [_source("Synthetic control assessment", "Existing controls represented: synthetic data boundary, human review, source traceability, and scoped charter. Design appears relevant; operating effectiveness has not been tested."), _source("Control gap", "Formal identity, access, retention, monitoring, cost, and incident-response controls require validation before institutional connection.")],
            "residual risk assessment": [_source("Demonstration residual assessment", "Residual probability 3/5; impact 4/5; residual rating 12/25 (High). Reduction reflects demo safeguards only; institutional tolerance and acceptance authority are not documented."), risk],
            "planned treatment view": (actions[:3] or evidence["recommendations"][:3]) + [_source("Target-state estimate", "Target rating 8/25 (Medium) after validated controls; target is prospective and must not be reported as current residual risk.")],
            "trend analysis": [_source("Trend limitation", "No prior approved risk scores exist, so numeric trend cannot be calculated."), _source("Qualitative trend", "Exposure is stable-to-improving if documented readiness actions close on time; otherwise it worsens as decision dates approach.")],
            "risk-priority view": [risk, _source("Priority assessment", "Executive review is warranted because residual exposure remains high and control effectiveness is unverified.")],
            "scoring-quality report": [_source("Scoring-quality assessment", "Missing: approved scoring framework, tolerance, control tests, prior score, specialist validation, and acceptance authority."), _source("Human-review control", "A qualified risk owner must validate every score and rationale; planned controls must remain separate from existing controls.")],
        }
        return mapping.get(value)

    if skill == "decision governance":
        mapping = {
            "decision-governance summary": (upcoming[:2] or decisions[:2]) + [_source("Governance assessment", "Decision level and formal authority are not fully documented; readiness is partial and specialist review may be required.")],
            "decision-rights matrix": [_source("Provisional rights matrix", "Requester: initiative owner | Preparer: program lead | Consultees: finance, operations, architecture, security/privacy | Approver and risk-acceptance authority: not documented | Informed: initiative participants."), _source("Authority limitation", "Roles are provisional and must not be treated as approved RACI or delegated authority.")],
            "governance pathway": [_source("Provisional governance pathway", "Request → classify → prepare evidence → consult specialists → compare options → authorized deliberation → decision → conditions → implementation → monitoring → review/closure."), _source("Pathway gap", "Authorized bodies, entry/exit criteria, target dates, escalation route, and source of authority require confirmation.")],
            "decision-readiness assessment": evidence["readiness"][:4] + [_source("Readiness classification", "Partially ready — core context exists, but authority, comparable options, cost/resource evidence, specialized reviews, and monitoring details remain incomplete.")],
            "option comparison": (evidence["tradeoffs"][:3] or evidence["recommendations"][:3]) + [_source("Comparison limitation", "Comparable option cost, feasibility, reversibility, stakeholder impact, and governance implications require validation.")],
            "deliberation brief": summary[:3] + evidence["questions"][:3],
            "decision record": evidence["decision_log_entries"][:4] or [_source("Decision record gap", "No completed decision record is available for the upcoming decision; create one after authorization.")],
            "conditions and obligations": actions[:4] + [_source("Governance condition", "Document evidence, monitoring, escalation triggers, and closure authority for every approval condition.")],
            "implementation and accountability plan": evidence["ownership"][:5],
            "decision monitoring": evidence["status"][:4] + [_source("Monitoring requirement", "Confirm outcome measures, review date, continued validity, and modification or closure authority.")],
            "governance gaps and conflicts": [_source("Governance gap", "Approving authority, risk-acceptance authority, and source of delegated authority are not represented."), _source("Governance gap", "Quorum, specialized review completion, implementation authorization, and monitoring cadence are not documented.")],
            "executive attention": (upcoming[:3] or risks[:3]) + [_source("Escalation requirement", "Escalate unresolved authority, material risk, time-sensitive readiness, or cross-functional disagreement to the authorized body.")],
            "source and confidence report": evidence["sources"] + [_source("Confidence assessment", "Medium-low: synthetic source records are traceable, but formal authority and institutional governance records are absent.")],
        }
        return mapping.get(value)

    if skill == "risk mitigation recommendations":
        mapping = {
            "risk summary": risks[:4] + [_source("Risk assessment limitation", "Inherent and residual numeric ratings require an approved framework and qualified owner validation.")],
            "mitigation options": [_source("Mitigation option", "Option 1 — Reduce: complete documented security, architecture, ownership, and dependency actions. Expected effect: lower probability; confidence: medium."), _source("Mitigation option", "Option 2 — Avoid: pause institutional connection until required controls are validated. Expected effect: lowers immediate exposure but delays value; confidence: high."), _source("Mitigation option", "Option 3 — Accept with monitoring: continue only within the synthetic, human-reviewed boundary. Expected effect: contains scope but does not resolve production readiness; confidence: medium.")],
            "comparative analysis": [_source("Mitigation comparison", "Reduce | Risk reduction: medium-high | Time: medium | Resource demand: medium-high | Reversible: yes"), _source("Mitigation comparison", "Avoid | Risk reduction: high | Time to value: delayed | Resource demand: low | Reversible: yes"), _source("Mitigation comparison", "Accept with monitoring | Risk reduction: low-medium | Time: immediate | Resource demand: low | Residual uncertainty: high")],
            "recommended mitigation": (actions[:3] or evidence["recommendations"][:3]) + [_source("Draft recommendation", "Use a phased Reduce strategy while retaining the synthetic-data and human-review boundary; obtain formal acceptance for any residual exposure.")],
            "contingency plan": [_source("Draft contingency", "Trigger: failed control, unapproved data access, material evidence contradiction, or missed readiness condition. Contain by stopping the affected activity, preserving records, notifying the owner, and escalating for review.")],
            "implementation roadmap": actions[:5] + evidence["schedule"][:3],
            "monitoring plan": [_source("Draft monitoring plan", "Track overdue actions, readiness dependencies, control validation, evidence contradictions, and decision dates; review at each linked governance meeting."), _source("Escalation threshold", "Escalate any unapproved connection, missing accountable owner, failed required control, or high residual risk.")],
            "executive decision request": (upcoming[:3] or evidence["approvals"][:3]) + [_source("Decision request", "Approve the mitigation direction, owner, resources, target dates, and residual-risk acceptance authority.")],
            "data-quality and validation summary": evidence["gaps"][:3] + [_source("Validation requirement", "Control effectiveness, cost, impact, residual risk, and specialist reviews require confirmation.")],
        }
        return mapping.get(value)

    return None


def _evidence_for_heading(heading: str, evidence: dict, skill_name: str = "") -> list[dict]:
    value = heading.lower()
    normalized_skill = skill_name.lower()
    specialized = _skill_specific_evidence(skill_name, heading, evidence)
    if specialized is not None:
        return specialized[:8]
    if value == "raci matrix" and normalized_skill == "stakeholder raci":
        return evidence["raci_matrix"]
    if value == "role summaries" and normalized_skill == "stakeholder raci":
        return evidence["raci_role_summaries"]
    if value == "executive summary" and normalized_skill == "stakeholder raci":
        return evidence["raci_executive_summary"][:6]
    if any(term in value for term in ("human-review requirement", "human review requirement")):
        if normalized_skill == "decision logs":
            return evidence["decision_human_review"][:6]
        if normalized_skill == "okr alignment":
            return evidence["okr_human_review"][:6]
    rules = (
        (("raci matrix",), "raci_matrix"),
        (("role summaries",), "raci_role_summaries"),
        (("responsibility gaps",), "raci_gaps"),
        (("responsibility conflicts",), "raci_conflicts"),
        (("workload and concentration",), "raci_workload"),
        (("governance and escalation",), "raci_governance"),
        (("executive okr summary",), "okr_summary"),
        (("okr inventory",), "okr_inventory"),
        (("objective quality assessment",), "objective_quality"),
        (("key-result quality assessment", "key result quality assessment"), "key_result_quality"),
        (("vertical alignment",), "vertical_alignment"),
        (("horizontal alignment",), "horizontal_alignment"),
        (("strategic coverage gap",), "okr_coverage_gaps"),
        (("recommended okr revision",), "okr_revisions"),
        (("progress and confidence assessment",), "okr_progress"),
        (("decision-log summary", "decision log summary"), "decision_summary"),
        (("decision-log entry", "decision log entry"), "decision_log_entries"),
        (("record completeness", "completeness status"), "record_completeness"),
        (("source", "traceability", "evidence"), "sources"),
        (("human review", "governance", "guardrail"), "governance"),
        (("success measure", "measure of success", "completion criteria", "acceptance criteria"), "success_measures"),
        (("capacity", "workload", "feasibility", "resource demand", "resource plan"), "capacity"),
        (("trade-off", "tradeoff", "defer", "delegate", "stop recommendation"), "tradeoffs"),
        (("priority rationale", "ranking rationale", "prioritization rationale"), "priority_rationale"),
        (("dependency map", "dependency schedule", "sequencing report", "sequence map"), "dependency_map"),
        (("readiness", "exception", "validation", "quality assessment"), "readiness"),
        (("schedule", "day-by-day", "five-day", "calendar of", "due by date", "look-ahead"), "schedule"),
        (("owner", "ownership", "responsibility", "raci", "accountability"), "ownership"),
        (("approval", "authorization", "disposition"), "approvals"),
        (("weekly outcome", "daily outcome", "outcome plan", "expected outcome"), "outcomes"),
        (("priority", "prioritized", "ranking", "ranked"), "priorities"),
        (("information gap", "confidence", "limitation"), "gaps"),
        (("participant", "stakeholder"), "participants"),
        (("meeting", "agenda", "calendar"), "meetings"),
        (("prior discussion", "history", "timeline", "what led"), "history"),
        (("upcoming decision", "leadership decision", "decisions required", "decision expected"), "upcoming_decisions"),
        (("decision", "alignment required", "decision register", "decision log"), "decisions"),
        (("action", "commitment", "follow-up", "implementation"), "actions"),
        (("risk", "issue", "dependency", "conflict"), "risks"),
        (("deliverable", "milestone", "work list"), "deliverables"),
        (("objective", "strategic", "okr", "portfolio", "coverage", "mapping"), "alignment"),
        (("question",), "questions"),
        (("recommend", "next step", "preparation", "resolution", "response"), "recommendations"),
        (("current status", "status dashboard", "current state"), "status"),
        (("overview", "profile", "context", "summary", "why"), "summary"),
    )
    for terms, key in rules:
        if any(_matches_heading(value, term) for term in terms):
            return evidence[key][:6]
    return evidence["overview"][:4]


def assess_inputs(inputs_markdown: str, evidence: dict) -> list[dict]:
    """Match documented input requirements to the available synthetic evidence."""
    checks = []
    available_text = " ".join(item["text"].lower() for values in evidence.values() for item in values)
    for requirement in _bullets(inputs_markdown)[:18]:
        tokens = [t for t in re.findall(r"[a-z]{4,}", requirement.lower()) if t not in {"required", "current", "relevant", "including", "information", "supporting"}]
        matched = any(token in available_text for token in tokens[:6])
        checks.append({"requirement": requirement, "status": "Available" if matched else "Not represented in synthetic demo"})
    return checks


def run_repository_skill(selected_skill: dict, context, scenario_name: str) -> dict:
    """Execute a skill contract against synthetic initiative evidence."""
    documents = selected_skill.get("documents", {})
    missing = [name for name in REQUIRED_DOCUMENTS if not documents.get(name, "").strip()]
    if missing:
        raise ValueError("Missing required repository files: " + ", ".join(missing))

    evidence = _initiative_evidence(context, scenario_name)
    sections = []
    for heading in _headings(documents["OUTPUTS.md"]):
        sections.append({"heading": heading, "items": _evidence_for_heading(heading, evidence, selected_skill["name"])})

    return {
        "skill_name": selected_skill["name"],
        "capability_name": selected_skill["capability"],
        "repository_path": selected_skill["repository_path"],
        "scenario_name": scenario_name,
        "sections": sections,
        "input_checks": assess_inputs(documents["INPUTS.md"], evidence),
        "documents_used": list(REQUIRED_DOCUMENTS),
        "review_status": "Human review required",
        "engine": "Repository-driven deterministic synthesis",
    }


def result_markdown(result: dict) -> str:
    lines = [
        f"# {result['skill_name']}",
        "",
        f"**Capability:** {result['capability_name']}  ",
        f"**Initiative:** {result['scenario_name']}  ",
        f"**Repository path:** `{result['repository_path']}`  ",
        f"**Review status:** {result['review_status']}",
        "",
    ]
    for section in result["sections"]:
        lines.extend([f"## {section['heading']}", ""])
        for item in section["items"]:
            lines.append(f"- {item['text']} **[{item['source']}]**")
        lines.append("")
    lines.extend([
        "## Governance notice",
        "",
        "Generated from repository skill contracts and synthetic demonstration records. No UW production systems were accessed. Human review is required.",
    ])
    return "\n".join(lines)


def result_docx(result: dict) -> bytes:
    doc = Document()
    doc.add_heading(result["skill_name"], 0)
    doc.add_paragraph(f"Capability: {result['capability_name']}")
    doc.add_paragraph(f"Initiative: {result['scenario_name']}")
    doc.add_paragraph(f"Repository path: {result['repository_path']}")
    doc.add_paragraph(f"Review status: {result['review_status']}")
    for section in result["sections"]:
        doc.add_heading(section["heading"], level=1)
        for item in section["items"]:
            doc.add_paragraph(f"{item['text']} [{item['source']}]", style="List Bullet")
    doc.add_heading("Governance notice", level=1)
    doc.add_paragraph("Generated from repository skill contracts and synthetic demonstration records. No UW production systems were accessed. Human review is required.")
    output = BytesIO()
    doc.save(output)
    return output.getvalue()
