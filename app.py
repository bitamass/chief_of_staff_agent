import re
from datetime import date
from io import BytesIO
from pathlib import Path

import streamlit as st
from docx import Document
from skill_catalog import discover_skill_catalog
from skill_runner import (
    REQUIRED_DOCUMENTS,
    result_docx,
    result_markdown,
    run_repository_skill,
)
from synthetic_context_v4 import DATASET_VERSION, build_initiative_scenarios, build_scenario_records, build_synthetic_context, initiative_continuity


st.set_page_config(
    page_title="Chief of Staff Agent",
    page_icon="🧭",
    layout="wide",
)

st.markdown(
    """
    <style>
    :root { --uw-purple:#4b2e83; --uw-gold:#b7a57a; --ink:#172033; --soft:#f5f3f8; }
    .block-container {padding-top: 2.2rem; max-width: 1280px;}
    h1, h2, h3 {color: var(--ink);}
    .hero {border-left:6px solid var(--uw-gold); padding:.25rem 0 .25rem 1rem; margin-bottom:1rem;}
    .hero h1 {margin:0; font-size:2.25rem;}
    .hero p {margin:.35rem 0 0; color:#5a6475;}
    .safe-banner {background:#f4f0fa; border:1px solid #d9cfea; border-radius:10px; padding:.8rem 1rem; margin:.7rem 0 1.2rem;}
    .brief-card {background:white; border:1px solid #e3e6eb; border-radius:12px; padding:1rem 1.1rem; min-height:145px; box-shadow:0 2px 8px rgba(20,30,50,.04);}
    .brief-card h4 {color:var(--uw-purple); margin:0 0 .5rem;}
    .source {color:#5b6677; font-size:.86rem;}
    div.stButton > button {background:var(--uw-purple); color:white; border:0; border-radius:8px; font-weight:650;}
    div.stButton > button:hover {background:#382261; color:white;}
    [data-testid="stMetric"] {background:#faf9fc; border:1px solid #e2dbea; padding:.8rem; border-radius:10px;}
    </style>
    """,
    unsafe_allow_html=True,
)

def find_sentences(records, keywords, limit=4, labels=None, excluded=None):
    found = []
    excluded = [x.lower() for x in (excluded or [])]
    for record in records:
        if labels and record["label"] not in labels:
            continue
        # Preserve paragraph boundaries so document titles and section labels do
        # not become fused with the first substantive sentence.
        sentences = re.split(r"\n+|(?<=[.!?])\s+|\s*[•]\s*", record["text"])
        for sentence in sentences:
            sentence = re.sub(r"^\s*\d+[.)]\s*", "", sentence).strip(" -\t")
            lowered = sentence.lower().rstrip(":")
            if lowered in {"agenda", "prior discussion notes", "background memo — executive meeting intelligence pilot"}:
                continue
            if "sample agenda" in lowered or lowered.startswith("background memo"):
                continue
            if any(term in lowered for term in excluded):
                continue
            for prefix in ("objective:", "decision requested:"):
                if lowered.startswith(prefix):
                    sentence = sentence.split(":", 1)[1].strip()
                    lowered = sentence.lower()
            if len(sentence) < 28:
                continue
            if any(term in sentence.lower() for term in excluded):
                continue
            if any(word.lower() in sentence.lower() for word in keywords):
                item = (sentence, record["label"])
                if item not in found:
                    found.append(item)
            if len(found) >= limit:
                return found
    return found


def make_demo_brief(title, objective, attendees, records, scenario_name=None):
    commitments = find_sentences(
        records,
        ["interest", "agreed", "requested", "supported", "required"],
        3,
        labels={"Prior Discussion Notes"},
        excluded=["no final decision"],
    )
    risks = find_sentences(
        records,
        ["concern", "risk", "permission", "access", "privacy", "security", "dependency", "downtime", "incident"],
        4,
        labels={"Prior Discussion Notes", "Background Memo"},
        excluded=["low-risk prototype", "no final decision", "suggested readiness measures"],
    )
    decisions = find_sentences(
        records,
        ["decision requested", "decide", "determine", "approve", "confirm", "identify"],
        3,
        labels={"Agenda"},
        excluded=["objective:", "no approval is requested"],
    )
    background = find_sentences(
        records,
        ["agent", "governance", "prototype", "executive", "initiative", "technology", "operational", "clinical", "investment", "demonstration"],
        4,
        labels={"Background Memo"},
    )

    # Keep the executive brief concise by preventing the same source sentence
    # from appearing in more than one section.
    used_text = {text for section in (commitments, risks, decisions) for text, _ in section}
    background = [item for item in background if item[0] not in used_text]

    if not commitments:
        commitments = [("No prior commitments were identified in the selected demonstration documents.", "Prior Discussion Notes")]
    if not risks:
        risks = [("The pilot requires explicit access controls, source traceability, human review, and an accountable owner.", "Background Memo")]
    if not decisions:
        decisions = [("No approval decision is requested; identify the questions that merit further exploration.", "Agenda")]

    default_scenario = next(iter(SCENARIOS.values()))
    scenario = SCENARIOS.get(scenario_name, default_scenario)
    questions = scenario["questions"]
    next_steps = scenario["next_steps"]
    return {
        "title": title,
        "objective": objective,
        "attendees": attendees,
        "background": background,
        "commitments": commitments,
        "risks": risks,
        "decisions": decisions,
        "questions": questions,
        "next_steps": next_steps,
    }


def cited_markdown(items):
    return "\n".join(f"- {text} **[{source}]**" for text, source in items)


def brief_markdown(brief, records):
    questions = "\n".join(f"- {x}" for x in brief["questions"])
    next_steps = "\n".join(f"- {x}" for x in brief["next_steps"])
    sources = "\n".join(f"- {r['label']} ({r['name']})" for r in records)
    return f"""# Executive Meeting Brief

**Meeting:** {brief['title']}  
**Objective:** {brief['objective']}  
**Attendees / roles:** {brief['attendees']}

## Relevant Background
{cited_markdown(brief['background'])}

## Prior Decisions and Commitments
{cited_markdown(brief['commitments'])}

## Open Issues and Risks
{cited_markdown(brief['risks'])}

## Decisions or Alignment Needed
{cited_markdown(brief['decisions'])}

## Recommended Questions
{questions}

## Recommended Next Steps
{next_steps}

## Sources Reviewed
{sources}

## Governance Notice
Generated from selected demonstration documents. Human review is required. This prototype is not connected to UW production systems.
"""


def brief_docx(brief, records):
    doc = Document()
    doc.add_heading("Executive Meeting Brief", 0)
    doc.add_paragraph(f"Meeting: {brief['title']}")
    doc.add_paragraph(f"Objective: {brief['objective']}")
    doc.add_paragraph(f"Attendees / roles: {brief['attendees']}")
    sections = [
        ("Relevant Background", brief["background"], True),
        ("Prior Decisions and Commitments", brief["commitments"], True),
        ("Open Issues and Risks", brief["risks"], True),
        ("Decisions or Alignment Needed", brief["decisions"], True),
        ("Recommended Questions", brief["questions"], False),
        ("Recommended Next Steps", brief["next_steps"], False),
    ]
    for heading, items, cited in sections:
        doc.add_heading(heading, level=1)
        for item in items:
            text = f"{item[0]} [{item[1]}]" if cited else item
            doc.add_paragraph(text, style="List Bullet")
    doc.add_heading("Sources Reviewed", level=1)
    for record in records:
        doc.add_paragraph(f"{record['label']} ({record['name']})", style="List Bullet")
    doc.add_heading("Governance Notice", level=1)
    doc.add_paragraph("Generated from selected demonstration documents. Human review is required. This prototype is not connected to UW production systems.")
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


st.markdown(
    '<div class="hero"><h1>Chief of Staff Agent</h1>'
    '<p>Exploratory demonstration of one executive-support skill.</p></div>',
    unsafe_allow_html=True,
)
st.markdown(
    '<div class="safe-banner"><b>Chiefs of Staff Demo:</b> This demo is based on synthetic data generated for this demo only. '
    'There are no connections to any UW production systems. Synthetic data, the demo App, and Chief of Staff skills '
    'are available for review in the GitHub repository.</div>',
    unsafe_allow_html=True,
)
st.caption(f"Synthetic dataset version: {DATASET_VERSION}")

operating_context = build_synthetic_context(date.today())
SCENARIOS = build_initiative_scenarios(operating_context)

st.subheader("Choose an initiative")
scenario_name = st.selectbox(
    "Initiative",
    list(SCENARIOS),
    label_visibility="collapsed",
    key="scenario_name",
)
scenario = SCENARIOS[scenario_name]
continuity = initiative_continuity(operating_context, scenario["initiative_id"])

st.markdown("#### People connected to this initiative")
people_columns = st.columns(min(len(continuity["participants"]), 4))
for index, person in enumerate(continuity["participants"]):
    with people_columns[index % len(people_columns)]:
        st.markdown(f"**{person['name']}**  \n{person['role']}")

capabilities = discover_capabilities(Path(__file__).parent)
capability_labels = list(capabilities)
default_capability_index = next((i for i, label in enumerate(capability_labels) if label.startswith("02 ")), 0)
st.markdown("#### Choose a Chief of Staff capability and skill")
capability_name = st.selectbox(
    "Capability",
    capability_labels,
    index=default_capability_index,
    key="capability_name",
)
skill_options = capabilities[capability_name]
preferred_skill = "Prepare Pre-Meeting Briefing Materials"
default_skill_index = skill_options.index(preferred_skill) if preferred_skill in skill_options else 0
skill_name = st.selectbox(
    "Skill",
    skill_options,
    index=default_skill_index,
    key=f"skill_name_{capability_name}",
)
if skill_name in CONNECTED_SKILLS:
    st.success(f"Selected repository skill: {capability_name} → {skill_name}")
else:
    st.info(
        f"{skill_name} is available in the repository framework but is not yet connected to this demo's synthetic inputs. "
        "Its SKILL, INPUTS, WORKFLOW, OUTPUTS, PROMPT, and EVALUATION files should be reviewed before activation."
    )

st.subheader("Executive timeline")
st.caption(f"Filtered to {scenario_name}: three weeks of operating history and two weeks of forward plans.")

st.markdown("#### Prior context")
past1, past2, past3 = st.columns(3)
past1.metric("Prior meetings", len(continuity["prior_meetings"]))
past2.metric("Prior decisions", len(continuity["prior_decisions"]))
past3.metric("Prior actions", len(continuity["prior_actions"]))

st.markdown("#### Forward view")
future1, future2, future3 = st.columns(3)
future1.metric("Upcoming meetings", len(continuity["upcoming_meetings"]))
future2.metric("Actions due and owners", len(continuity["actions_due"]))
future3.metric("Upcoming decisions", len(continuity["upcoming_decisions"]))

st.markdown("#### Meetings, decisions, and actions for the selected initiative")
detail_meetings, detail_decisions, detail_actions = st.columns(3)
with detail_meetings:
    st.markdown("##### Meetings")
    with st.container(border=True):
        st.markdown("**Prior**")
        if continuity["prior_meetings"]:
            for item in reversed(continuity["prior_meetings"]):
                st.markdown(f"- **{item['date']:%b %d}:** {item['title']}")
        else:
            st.caption("No prior meetings recorded.")
    with st.container(border=True):
        st.markdown("**Current and upcoming**")
        for item in continuity["current_meetings"]:
            st.markdown(f"- **Today · {item['date']:%b %d}:** {item['title']}")
        for item in continuity["upcoming_meetings"]:
            st.markdown(f"- **{item['date']:%b %d}:** {item['title']}")
with detail_decisions:
    st.markdown("##### Decisions")
    with st.container(border=True):
        st.markdown("**Prior**")
        if continuity["prior_decisions"]:
            for item in reversed(continuity["prior_decisions"]):
                source_meeting = operating_context["meetings"][item["meeting_id"]]
                st.markdown(f"- **{item['date']:%b %d}:** {item['decision']}  \n  <span class='source'>{source_meeting['title']}</span>", unsafe_allow_html=True)
        else:
            st.caption("No prior decisions recorded.")
    with st.container(border=True):
        st.markdown("**Upcoming**")
        if continuity["upcoming_decisions"]:
            for item in continuity["upcoming_decisions"]:
                meeting = operating_context["meetings"][item["meeting_id"]]
                st.markdown(f"- **{item['date']:%b %d}:** {item['decision']}  \n  <span class='source'>{meeting['title']}</span>", unsafe_allow_html=True)
        else:
            st.caption("No upcoming decisions recorded.")
with detail_actions:
    st.markdown("##### Actions")
    with st.container(border=True):
        st.markdown("**Prior**")
        if continuity["prior_actions"]:
            for item in continuity["prior_actions"]:
                owner = operating_context["participants"][item["owner"]]
                st.markdown(f"- **{item['status']} · Due {item['due_date']:%b %d}:** {item['action']}  \n  <span class='source'>Owner: {owner['name']} — {owner['role']}</span>", unsafe_allow_html=True)
        else:
            st.caption("No prior actions recorded.")
    with st.container(border=True):
        st.markdown("**Upcoming**")
        if continuity["actions_due"]:
            for item in continuity["actions_due"]:
                owner = operating_context["participants"][item["owner"]]
                st.markdown(f"- **{item['status']} · Due {item['due_date']:%b %d}:** {item['action']}  \n  <span class='source'>Owner: {owner['name']} — {owner['role']}</span>", unsafe_allow_html=True)
        else:
            st.caption("No upcoming actions recorded.")

st.markdown("#### Selected Chief of Staff skill output")
skill_records = build_scenario_records(operating_context, scenario_name)
if skill_name in {"Action Item Management", "Follow Up On Commitments"}:
    for item in continuity["all_actions"]:
        owner = operating_context["participants"][item["owner"]]
        st.markdown(f"- **{item['status']} · Due {item['due_date']:%b %d}:** {item['action']} — {owner['name']}, {owner['role']}")
elif skill_name == "Calendar Prioritization":
    for priority, item in enumerate(continuity["current_meetings"] + continuity["upcoming_meetings"], start=1):
        st.markdown(f"- **Priority {priority} · {item['date']:%b %d}:** {item['title']} — {item['summary']}")
elif skill_name == "Create Agendas":
    st.code(skill_records[0]["text"], language="text")
elif skill_name == "Meeting Preparation":
    st.markdown(f"**Meeting:** {scenario['title']}  \n**Date:** {scenario['date']:%b %d, %Y}  \n**Objective:** {scenario['objective']}  \n**Attendees:** {scenario['attendees']}")
elif skill_name == "Prepare Pre-Meeting Briefing Materials":
    st.write("The meeting context and three source-linked briefing materials below are ready for executive brief generation.")
elif skill_name == "Summarize Prior Discussions":
    for item in reversed(continuity["prior_meetings"]):
        st.markdown(f"- **{item['date']:%b %d} · {item['title']}:** {item['summary']}")
elif skill_name in {"Track Decisions", "Decision Logs"}:
    for item in reversed(continuity["prior_decisions"]):
        st.markdown(f"- **Made {item['date']:%b %d}:** {item['decision']}")
    for item in continuity["upcoming_decisions"]:
        st.markdown(f"- **Upcoming {item['date']:%b %d}:** {item['decision']}")
elif skill_name == "Initiative Portfolio Review":
    initiative = operating_context["initiatives"][scenario["initiative_id"]]
    charter = operating_context["charters"][initiative["charter_id"]]
    st.markdown(f"**Status:** {initiative['status']}  \n**Objective:** {initiative['objective']}  \n**Charter:** {charter['title']}  \n**Success measures:** {', '.join(charter['success_measures'])}")
    for item in operating_context["deliverables"].values():
        if item["initiative_id"] == scenario["initiative_id"]:
            st.markdown(f"- **{item['status']} · Due {item['due_date']:%b %d}:** {item['name']}")
else:
    st.warning("This skill is listed for framework review only. It is not executed because its repository instructions and required synthetic inputs have not yet been connected.")

with st.expander("Explore the complete five-week synthetic calendar and supporting records"):
    calendar_rows = []
    for day in operating_context["calendar"]:
        event = day["meeting"]
        calendar_rows.append({
            "Date": day["date"].strftime("%a, %b %d"),
            "Relative day": "Today" if day["offset"] == 0 else f"{day['offset']:+d}",
            "Meeting": event["title"] if event else "—",
            "Context": event["summary"] if event else "No scheduled executive meeting; continuity retained.",
        })
    st.dataframe(calendar_rows, width="stretch", hide_index=True)
    st.markdown("#### Documented backend records")
    st.write("**Participants:** " + ", ".join(p["role"] for p in operating_context["participants"].values()))
    st.write("**Initiatives:** " + "; ".join(i["name"] for i in operating_context["initiatives"].values()))
    st.write("**Charters:** " + "; ".join(c["title"] for c in operating_context["charters"].values()))
    st.write("**Deliverables:** " + "; ".join(d["name"] for d in operating_context["deliverables"].values()))
    st.caption("Every meeting, decision, action, initiative, charter, and deliverable uses a synthetic identifier so related records can be traced across time.")

with st.sidebar:
    st.header("Demonstration Safeguards")
    st.success("Chiefs of Staff Demo")
    for label in [
        "Human review required",
        "Source traceability enabled",
        "Session-only document use",
        "No production or Epic data",
    ]:
        st.checkbox(label, value=True, disabled=True)
    st.divider()
    st.subheader("Future state—not connected")
    st.caption("Outlook calendar · Teams summaries · SharePoint · Approved external intelligence")

st.subheader("Most relevant upcoming meeting")
scenario_action, scenario_description = st.columns([1, 2.4])
with scenario_action:
    if st.button("Update Relevant Meeting Materials", width="stretch"):
        st.session_state.meeting_title = scenario["title"]
        st.session_state.meeting_objective = scenario["objective"]
        st.session_state.attendees = scenario["attendees"]
        st.session_state.meeting_date = scenario["date"]
        st.session_state.active_scenario = scenario_name
        st.session_state.pop("brief", None)
        st.rerun()
with scenario_description:
    st.caption("Updates the agenda, prior notes, background information, and meeting context for the selected initiative.")

default_scenario = next(iter(SCENARIOS.values()))
if "meeting_title" not in st.session_state:
    st.session_state.meeting_title = default_scenario["title"]
if "meeting_objective" not in st.session_state:
    st.session_state.meeting_objective = default_scenario["objective"]
if "attendees" not in st.session_state:
    st.session_state.attendees = default_scenario["attendees"]
if "meeting_date" not in st.session_state:
    st.session_state.meeting_date = default_scenario["date"]

left, right = st.columns([1.15, 1])
with left:
    meeting_title = st.text_input("Meeting title", key="meeting_title")
    meeting_objective = st.text_area(
        "Meeting objective",
        key="meeting_objective",
        height=95,
    )
with right:
    meeting_date = st.date_input("Meeting date", key="meeting_date")
    attendees = st.text_area("Attendees / roles", key="attendees", height=95)

records = []
if st.session_state.get("active_scenario"):
    active_name = st.session_state.active_scenario
    records = build_scenario_records(operating_context, active_name)

if records:
    st.info("Sources ready: " + ", ".join(r["name"] for r in records))

if st.button("Create Executive Meeting Brief", width="stretch"):
    if not records:
        st.warning("Select Update Relevant Meeting Materials before creating the brief.")
    else:
        st.session_state.brief = make_demo_brief(
            meeting_title,
            meeting_objective,
            attendees,
            records,
            st.session_state.get("active_scenario"),
        )
        st.session_state.records = records
        st.session_state.generated_date = str(meeting_date)

if "brief" in st.session_state:
    brief = st.session_state.brief
    saved_records = st.session_state.records
    st.divider()
    st.header("Executive Brief")
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Sources reviewed", len(saved_records))
    m2.metric("Decisions surfaced", len(brief["decisions"]))
    m3.metric("Risks identified", len(brief["risks"]))
    m4.metric("Review status", "Human review")

    st.subheader(brief["title"])
    st.caption(f"Meeting date: {st.session_state.generated_date}  ·  Prepared from selected demonstration documents")
    st.write(brief["objective"])

    a, b = st.columns(2)
    with a:
        st.markdown("#### Decisions or Alignment Needed")
        for text, source in brief["decisions"]:
            st.markdown(f"- {text}  \n  <span class='source'>Source: {source}</span>", unsafe_allow_html=True)
    with b:
        st.markdown("#### Open Issues and Risks")
        for text, source in brief["risks"]:
            st.markdown(f"- {text}  \n  <span class='source'>Source: {source}</span>", unsafe_allow_html=True)

    st.markdown("#### Prior Decisions and Commitments")
    for text, source in brief["commitments"]:
        st.markdown(f"- {text} **[{source}]**")

    with st.expander("Relevant background and evidence", expanded=True):
        for text, source in brief["background"]:
            st.markdown(f"- {text} **[{source}]**")

    qcol, ncol = st.columns(2)
    with qcol:
        st.markdown("#### Recommended Questions")
        for item in brief["questions"]:
            st.markdown(f"- {item}")
    with ncol:
        st.markdown("#### Recommended Next Steps")
        for item in brief["next_steps"]:
            st.markdown(f"- {item}")

    with st.expander("Agent governance record"):
        st.write("**Purpose:** Apply the selected Chief of Staff skill to the selected synthetic initiative context.")
        st.write("**Data accessed:** " + ", ".join(r["name"] for r in saved_records))
        st.write("**External systems accessed:** None")
        st.write("**Retention:** Session only")
        st.write("**Required control:** Human review before use or distribution")
        st.write("**Known limitation:** Demonstration synthesis; no automated validation against UW systems")

    md = brief_markdown(brief, saved_records)
    docx = brief_docx(brief, saved_records)
    d1, d2 = st.columns(2)
    d1.download_button("Download Word Brief", docx, "executive_meeting_brief.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", width="stretch")
    d2.download_button("Download Markdown Brief", md, "executive_meeting_brief.md", "text/markdown", width="stretch")

with st.expander("Chief of Staff capability framework discovered from the repository"):
    for capability, skills in capabilities.items():
        st.markdown(f"**{capability}**")
        st.write(" · ".join(skills))
    st.caption("A skill is executed only when its repository instructions and required synthetic inputs are connected to this demo.")

st.divider()
st.caption("Prototype boundary: synthetic data only; no autonomous actions; no production connections. Any institutional pilot would require approved authentication, role-based access, logging, privacy, security, records-management, and governance review.")
